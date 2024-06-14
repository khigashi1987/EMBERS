import os
import io
import glob
import logging
import pandas as pd
import docx
from pypdf import PdfReader
from camelot import read_pdf
from config import Config

class SUPPMATLoader():
    ###
    # Extract texts and tables from Supplementally materials.
    ###
    def __init__(self):
        pass

    def columns_match(self, columns1, columns2):
        if len(columns1) != len(columns2):
            return False
        else:
            return True
    
    def extract_texts_docx(self, docx_file):
        doc = docx.Document(docx_file)
        texts = ''
        for para in doc.paragraphs:
            texts += para.text+'\n'
        return texts

    def extract_texts(self, pdf_file):
        pdf = PdfReader(pdf_file)
        n_pages = len(pdf.pages)
        texts = '' 
        for page in range(n_pages):
            texts += '\n'+pdf.pages[page].extract_text()+'\n'
        return texts

    def check_if_fix_needed(self, df):
        # データフレームを検証して、ヘッダの修正が必要かどうかを判断する
        # カラム名のリストで重複をチェック
        columns = df.columns.tolist()
        unique_columns = set(columns)
        if len(columns) != len(unique_columns):
            if all(isinstance(item, str) for item in df.iloc[0]):
                # 最初の行が実際のデータではなく、カラム名のような場合
                return True
        return False
    
    def fix_merged_headers(self, df):
        # ヘッダ行を修正するための関数
        # 最初の行を抽出し、カラム名と組み合わせる
        new_headers = df.iloc[0]
        df = df[1:]  # 最初の行を除外
        df.columns = [f"{col}.{nh}" if not col == nh else col for col, nh in zip(df.columns, new_headers)]
        return df

    def extract_tables(self, pdf_file):
        # get page number of PDF file
        pdf = PdfReader(pdf_file)
        n_pages = len(pdf.pages)
        last_columns = None
        last_page = None
        combined_tables = []

        # extract tables for each page
        for page in range(n_pages):
            tables = read_pdf(pdf_file, suppress_stdout=True,
                              flavor='stream', pages=str(page+1))
            if len(tables) == 0:
                continue

            for table in tables:
                current_df = table.df
                if current_df.shape[1] < 2:
                    continue
                if last_columns is not None and \
                    self.columns_match(last_columns, current_df.columns) and \
                    last_page + 1 == page:
                    # テーブルが複数ページにまたがっている場合、
                    # 前のテーブルに現在のテーブルを連結
                    combined_tables[-1] = pd.concat([combined_tables[-1], current_df], ignore_index=True)
                else:
                    # 新しいテーブルとして処理
                    combined_tables.append(current_df)
                
                last_columns = current_df.columns
                last_page = page  # 現在のページ番号を更新

        return combined_tables

    def extract_tables_docx(self, docx_file, header_row=True):
        doc = docx.Document(docx_file)
        tables = []
        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                data.append(row_data)
            if header_row:
                df = pd.DataFrame(data[1:], columns=data[0])
            else:
                df = pd.DataFrame(data)
            if self.check_if_fix_needed(df):
                df = self.fix_merged_headers(df)
            tables.append(df)
        return tables

    def analyze_suppmat(self, pmc_dir):
        article_nxml = glob.glob(os.path.join(pmc_dir, '*.nxml'))[0]
        article_id = article_nxml.split('/')[-1].split('.')[0]
        article_supp_pdf_list = glob.glob(os.path.join(pmc_dir, '*.pdf'))
        article_supp_docx_list = glob.glob(os.path.join(pmc_dir, '*.docx'))

        all_tables = []
        supp_method_text = ''
        for pdf_file in article_supp_pdf_list:
            # In Text extraction process, main article body should be skipped.
            if not os.path.basename(pdf_file).startswith(article_id):
                supp_method_text += self.extract_texts(pdf_file)
            
            # Skip if the file size is too large
            if os.path.getsize(pdf_file) > Config.SKIP_SUPP_SIZE:
                logging.info(f'\t\tSkip large file: {pdf_file}')
                continue

            # Table extraction process including tables main article body
            logging.info(f'\t\tanalyzing {pdf_file}')

            try:
                for t in self.extract_tables(pdf_file):
                    # 75%以上の列がNaNである行を削除する
                    threshold_row = len(t.columns) * 0.75
                    t = t.dropna(axis=0, thresh=threshold_row)
                    # 75%以上の行がNaNである列を削除する
                    threshold_col = len(t) * 0.75
                    t = t.dropna(axis=1, thresh=threshold_col)
                    # if columns are just numbers,
                    # set the first row as column names
                    columns_as_strings = t.columns.astype(str)
                    if all(column.isnumeric() for column in columns_as_strings):
                        new_header = t.iloc[0]
                        t = t.iloc[1:]
                        t.columns = new_header
                    # if columns contains NaN or empty string,
                    # set the first row as column names
                    if any(column == '' or column is None for column in t.columns):
                        new_header = t.iloc[0]
                        t = t.iloc[1:]
                        t.columns = new_header
                    all_tables.append(t)
            except Exception as e:
                print(e)
                logging.error(f'\t\tPDF Parsing Error: {e}')

            logging.info(f'\t\tanalyze done. {pdf_file}')
        
        for docx_file in article_supp_docx_list:
            # Skip if the file size is too large
            if os.path.getsize(docx_file) > Config.SKIP_SUPP_SIZE:
                logging.info(f'\t\tSkip large file: {docx_file}')
                continue

            logging.info(f'\t\tanalyzing {docx_file}')
            try:
                supp_method_text += self.extract_texts_docx(docx_file)
                extracted_tables_docx = self.extract_tables_docx(docx_file)
                if len(extracted_tables_docx) > 0:
                    all_tables += extracted_tables_docx
            except Exception as e:
                print(e)
                logging.error(f'\t\tDOCX Parsing Error: {e}')

            logging.info(f'\t\tanalyze done. {docx_file}')
        
        return supp_method_text, all_tables

if __name__ == '__main__':
    import sys
    if len(sys.argv) != 2:
        print(f'Usage: python {sys.argv[0]} <PDF file>')
        sys.exit(1)
    pdfloader = SUPPMATLoader()
    tables = pdfloader.extract_tables(sys.argv[1])
    for i, t in enumerate(tables):
        print(f'Table {i+1}')
        print(t)
        print('\n')

